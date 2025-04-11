from docx import Document
from io import BytesIO

def create_problem_and_answer_docs(text: str):
    problems = []
    answers = []

    in_answers = False
    for line in text.split("\n"):
        if line.strip().startswith("정답:"):
            in_answers = True
            answers.append(line.strip())
        elif in_answers:
            answers.append(line.strip())
        else:
            problems.append(line.strip())

    problem_doc = Document()
    for p in problems:
        problem_doc.add_paragraph(p)

    problem_io = BytesIO()
    problem_doc.save(problem_io)
    problem_io.seek(0)

    answer_doc = Document()
    for a in answers:
        answer_doc.add_paragraph(a)

    answer_io = BytesIO()
    answer_doc.save(answer_io)
    answer_io.seek(0)

    return problem_io, answer_io
