import fitz  # PyMuPDF
import re
import pandas as pd
from docx import Document
import io

def create_problem_and_answer_docs(gpt_output):
    parts = gpt_output.split("정답:", 1)
    problem_text = parts[0].strip()
    answer_text = parts[1].strip() if len(parts) > 1 else ""

    # 문제 문서 생성
    problem_doc = Document()
    problem_doc.add_heading('단어 문제', level=1)
    for line in problem_text.split("\n"):
        problem_doc.add_paragraph(line)

    # 정답 문서 생성
    answer_doc = Document()
    answer_doc.add_heading('정답', level=1)
    for line in answer_text.split("\n"):
        answer_doc.add_paragraph(line)

    # 메모리에 저장 후 BytesIO 객체 반환
    problem_io = io.BytesIO()
    answer_io = io.BytesIO()
    problem_doc.save(problem_io)
    answer_doc.save(answer_io)
    problem_io.seek(0)
    answer_io.seek(0)
    return problem_io.getvalue(), answer_io.getvalue()
    
def extract_units_from_excel(file):
    df = pd.read_excel(file)
    units = {}

    # "Unit"과 "Word", "Definition" 컬럼이 있다고 가정
    for _, row in df.iterrows():
        unit = row["Unit"]
        word = row["Word"]
        definition = row["Definition"]

        if pd.isna(unit) or pd.isna(word):
            continue

        unit_key = f"Unit {int(unit)}" if isinstance(unit, (int, float)) else str(unit)
        if unit_key not in units:
            units[unit_key] = []
        units[unit_key].append({
            "word": str(word).strip(),
            "definition": str(definition).strip() if pd.notna(definition) else ""
        })

    return units

def extract_units_individually_from_pdf(file):
    file.seek(0)
    doc = fitz.open(stream=file.read(), filetype="pdf")

    full_text = ""
    for page in doc:
        full_text += page.get_text("text") + "\n"

    units = {}
    current_unit = None
    current_words = []

    lines = full_text.split("\n")
    unit_pattern = re.compile(r"^Unit\s+(\d+)\b", re.IGNORECASE)
    word_line_pattern = re.compile(r"^([a-zA-Z\-']+)\s+\[[^\]]+\]")  # 단어 + 발음
    pos_definition_pattern = re.compile(r"^(n\.|v\.|adj\.|adv\.)\s+(.+)", re.IGNORECASE)

    last_word = None  # 직전 단어 저장용

    for line in lines:
        line = line.strip()

        if line.startswith("Review") or line.startswith("Practice"):
            continue

        unit_match = unit_pattern.match(line)
        if unit_match:
            if current_unit and current_words:
                units[f"Unit {current_unit}"] = current_words
            current_unit = unit_match.group(1)
            current_words = []
            continue

        word_match = word_line_pattern.match(line)
        if word_match:
            word = word_match.group(1)
            last_word = {"word": word, "definition": ""}
            current_words.append(last_word)
            continue

        # 영영 뜻 줄 감지: 품사로 시작하면 해당 단어에 연결
        definition_match = pos_definition_pattern.match(line)
        if definition_match and last_word:
            definition = definition_match.group(2).strip()
            last_word["definition"] = definition

    if current_unit and current_words:
        units[f"Unit {current_unit}"] = current_words

    # 출력 확인
    for unit, words in units.items():
        print(f"{unit}:\n{'-' * len(unit)}")
        for idx, entry in enumerate(words, 1):
            print(f"{idx}. {entry['word']} - {entry['definition']}")
        print("\n")

    return units

def extract_units_from_docx(file):
    doc = docx.Document(file)
    full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

    # 유닛 단위로 split (예: Unit 1 First Day)
    units = re.split(r'(Unit \d+[^U]*)', full_text)
    
    unit_data = {}
    for i in range(1, len(units), 2):
        title = units[i].strip().split("\n")[0]  # ex: Unit 1 First Day
        content = units[i] + units[i+1] if i + 1 < len(units) else units[i]
        unit_data[title] = content.strip()
    
    return unit_data


def separate_problems(text: str):
    problem_pattern = re.compile(r"(?P<number>\[\d+\s*~\s*\d+\]|\d{1,2})\.\s*(?P<question>.+?)(?=\n|$)")
    choice_pattern = re.compile(r"(①|②|③|④|⑤)")

    lines = text.strip().split("\n")
    problems = []
    current_problem = {}
    buffer = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        problem_match = problem_pattern.match(line)
        if problem_match:
            if current_problem:
                combined = "\n".join(buffer).strip()
                split_match = choice_pattern.search(combined)
                if split_match:
                    split_idx = split_match.start()
                    current_problem["passage"] = combined[:split_idx].strip()
                    current_problem["choices"] = combined[split_idx:].strip()
                else:
                    current_problem["passage"] = combined
                    current_problem["choices"] = ""
                problems.append(current_problem)
                buffer = []

            number = problem_match.group("number")
            question = problem_match.group("question")
            current_problem = {
                "number": number,
                "question": question,
                "passage": "",
                "choices": ""
            }
        else:
            buffer.append(line)

    if current_problem:
        combined = "\n".join(buffer).strip()
        split_match = choice_pattern.search(combined)
        if split_match:
            split_idx = split_match.start()
            current_problem["passage"] = combined[:split_idx].strip()
            current_problem["choices"] = combined[split_idx:].strip()
        else:
            current_problem["passage"] = combined
            current_problem["choices"] = ""
        problems.append(current_problem)

    return problems


def parse_primary_level_questions(text: str):
    pattern = re.compile(r"(?P<number>\d+)\.\s*(?:\n)?(?P<choices>(?:[a-d]\).+?\n?)+)", re.MULTILINE)
    choice_pattern = re.compile(r"(a|b|c|d)\)\s*(.+)")
    problems = []

    for match in pattern.finditer(text):
        number = match.group("number")
        raw_choices = match.group("choices").strip().split("\n")
        parsed_choices = []
        for c in raw_choices:
            choice_match = choice_pattern.match(c.strip())
            if choice_match:
                parsed_choices.append(choice_match.group(2))
        problems.append({
            "number": number,
            "question": "보기 중 알맞은 단어를 고르세요.",
            "choices": parsed_choices
        })
    return problems


def extract_text_from_pptx(file):
    prs = Presentation(file)
    text = ""
    for slide in prs.slides:
        for shape in slide.shapes:
            if hasattr(shape, "text"):
                text += shape.text + "\n"
    return text
